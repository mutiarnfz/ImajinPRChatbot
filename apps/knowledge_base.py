"""
Basis pengetahuan perusahaan, diambil dari file .docx yang ditaruh di
folder `knowledge/`.

Pendekatan: setiap file .docx dibaca langsung memakai library `python-docx`
(murni Python, tidak butuh aplikasi eksternal seperti pandoc terpasang di
sistem), lalu semua isinya (paragraf + tabel) digabung jadi satu konteks
teks yang disuntikkan ke system prompt saat menjawab pertanyaan di
endpoint /profil.

Cocok untuk dokumen berukuran kecil-menengah (company profile, FAQ, dsb).
Jika nanti dokumennya banyak/besar, pertimbangkan upgrade ke pendekatan
RAG (chunking + embedding + vector search) agar tidak melebihi context
window model.
"""

import os
import threading

from docx import Document

import config

_lock = threading.Lock()
_cache: str | None = None


def _convert_docx_to_text(path: str) -> str:
    """Ekstrak teks (paragraf & tabel) dari satu file .docx."""
    try:
        doc = Document(path)
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        return "\n".join(parts)
    except Exception as e:
        # Jangan hentikan seluruh aplikasi hanya karena satu file gagal dibaca
        print(f"[knowledge_base] Gagal membaca {path}: {e}")
        return ""


def _load_all() -> str:
    if not os.path.isdir(config.KNOWLEDGE_DIR):
        return ""

    chunks = []
    for filename in sorted(os.listdir(config.KNOWLEDGE_DIR)):
        if filename.lower().endswith(".docx"):
            full_path = os.path.join(config.KNOWLEDGE_DIR, filename)
            text = _convert_docx_to_text(full_path)
            if text.strip():
                chunks.append(f"## Sumber: {filename}\n\n{text.strip()}")

    return "\n\n---\n\n".join(chunks)


def get_company_knowledge(force_reload: bool = False) -> str:
    """Mengembalikan gabungan teks seluruh .docx di knowledge/.
    Di-cache di memori agar tidak baca ulang file di setiap request;
    panggil dengan force_reload=True untuk memuat ulang setelah file
    knowledge diperbarui."""
    global _cache
    with _lock:
        if _cache is None or force_reload:
            _cache = _load_all()
        return _cache
